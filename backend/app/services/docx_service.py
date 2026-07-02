"""Servicio de digitalización: convierte texto OCR en un .docx descargable.

El .docx puede construirse de dos formas:
  1. A partir de "bloques" estructurados (títulos, párrafos, viñetas, tablas),
     típicamente generados por Gemini a partir del texto OCR — reconstruye la
     estructura real del documento (incluidas tablas).
  2. A partir del texto OCR plano, aplicando un parser heurístico local (sin IA)
     que detecta títulos numerados y viñetas. Es el fallback gratuito.
"""
import io
import logging
import re
from datetime import datetime, timezone
from typing import Optional

logger = logging.getLogger("docmind")

# Marcadores de viñeta que suele producir el OCR ("+", "-", "*", "•", "»", …).
_BULLET_RE = re.compile(r"^\s*[-+*•·»▪◦‣]\s+(.*)")
# Título numerado de sección: "1. Resumen Ejecutivo", "5. Plan Q2 2026".
_HEADING_RE = re.compile(r"^\s*\d+\.\s+[A-Za-zÁÉÍÓÚÜÑáéíóúüñ].*")


def _heuristic_blocks(ocr_text: str) -> list[dict]:
    """
    Parser local (sin IA): convierte el texto OCR en bloques estructurados.

    - Líneas que empiezan con un marcador de viñeta → bloque 'bullets'.
    - Líneas con "N. Título" → bloque 'heading' (nivel 2).
    - El resto de líneas consecutivas se agrupan en párrafos.
    No reconstruye tablas (eso requiere IA o análisis de layout).
    """
    blocks: list[dict] = []
    para_buf: list[str] = []
    bullet_buf: list[str] = []

    def flush_para() -> None:
        if para_buf:
            blocks.append({"type": "paragraph", "text": " ".join(para_buf).strip()})
            para_buf.clear()

    def flush_bullets() -> None:
        if bullet_buf:
            blocks.append({"type": "bullets", "items": list(bullet_buf)})
            bullet_buf.clear()

    for raw in ocr_text.split("\n"):
        line = raw.strip()
        if not line:
            flush_para()
            flush_bullets()
            continue

        bullet = _BULLET_RE.match(line)
        if bullet:
            flush_para()
            bullet_buf.append(bullet.group(1).strip())
            continue

        if _HEADING_RE.match(line):
            flush_para()
            flush_bullets()
            blocks.append({"type": "heading", "level": 2, "text": line})
            continue

        flush_bullets()
        para_buf.append(line)

    flush_para()
    flush_bullets()
    return blocks


def _render_blocks(doc, blocks: list[dict]) -> None:
    """Renderiza una lista de bloques estructurados en el documento python-docx."""
    for block in blocks:
        btype = block.get("type")

        if btype == "heading":
            level = block.get("level", 2)
            level = max(1, min(int(level) if str(level).isdigit() else 2, 4))
            text = (block.get("text") or "").strip()
            if text:
                doc.add_heading(text, level=level)

        elif btype == "bullets":
            for item in block.get("items", []):
                item = (item or "").strip()
                if item:
                    doc.add_paragraph(item, style="List Bullet")

        elif btype == "table":
            _render_table(doc, block)

        else:  # "paragraph" o desconocido
            text = (block.get("text") or "").strip()
            if text:
                doc.add_paragraph(text)


def _render_table(doc, block: dict) -> None:
    """Renderiza un bloque de tabla; si viene malformado, cae a párrafo."""
    header = block.get("header") or []
    rows = block.get("rows") or []
    if not rows and not header:
        return

    n_cols = max(
        [len(header)] + [len(r) for r in rows if isinstance(r, list)] or [0]
    )
    if n_cols == 0:
        return

    table = doc.add_table(rows=0, cols=n_cols)
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        table.style = "Table Grid"

    def _add_row(values: list, bold: bool = False) -> None:
        cells = table.add_row().cells
        for i in range(n_cols):
            val = str(values[i]).strip() if i < len(values) and values[i] is not None else ""
            cells[i].text = val
            if bold:
                for para in cells[i].paragraphs:
                    for run in para.runs:
                        run.bold = True

    if header:
        _add_row(header, bold=True)
    for row in rows:
        if isinstance(row, list):
            _add_row(row)


def build_docx(
    ocr_text: str,
    source_filename: str,
    blocks: Optional[list[dict]] = None,
) -> bytes:
    """
    Genera un documento .docx editable a partir del texto OCR.

    Args:
        ocr_text:        Texto OCR del documento (fallback si no hay `blocks`).
        source_filename: Nombre del archivo original (para el encabezado).
        blocks:          Bloques estructurados (Gemini). Si es None, se usa el
                         parser heurístico local sobre `ocr_text`.

    Incluye un encabezado con el nombre del archivo original y la fecha de
    digitalización, seguido por el contenido estructurado. Retorna los bytes.
    """
    from docx import Document  # noqa: PLC0415
    from docx.shared import Pt  # noqa: PLC0415

    doc = Document()

    # Estilo base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Encabezado
    title = doc.add_heading("Documento digitalizado", level=1)
    title.alignment = 0

    meta = doc.add_paragraph()
    meta.add_run("Archivo original: ").bold = True
    meta.add_run(source_filename)
    meta_date = doc.add_paragraph()
    meta_date.add_run("Digitalizado el: ").bold = True
    meta_date.add_run(datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC"))

    doc.add_paragraph()  # espacio en blanco

    text = (ocr_text or "").strip()
    render_blocks = blocks if blocks else (_heuristic_blocks(text) if text else [])

    if not render_blocks:
        doc.add_paragraph(
            "[No se pudo extraer texto del documento original. "
            "Revise manualmente la imagen o el escaneo.]"
        )
    else:
        _render_blocks(doc, render_blocks)

    buffer = io.BytesIO()
    doc.save(buffer)
    data = buffer.getvalue()
    buffer.close()

    logger.info(
        f"DOCX generado: {len(data)} bytes "
        f"({len(text)} chars de OCR, {len(render_blocks)} bloques, "
        f"estructura={'gemini' if blocks else 'heuristica'}, "
        f"archivo origen={source_filename!r})"
    )
    return data
