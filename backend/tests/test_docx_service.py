"""Tests del servicio de digitalización a .docx.

Cubren el parser heurístico (Python puro) y la generación del .docx tanto desde
texto plano como desde bloques estructurados (incluida tabla). Pillow/python-docx
están disponibles en el entorno de tests.
"""
import io

from app.services import docx_service


OCR_SAMPLE = """\
1. Resumen Ejecutivo del Trimestre
El presente reporte consolida el avance de los proyectos activos de Nexus
Consulting S.A.C. durante el primer trimestre del 2026.
3. Métricas del Trimestre
+ Total de horas facturadas: 1,840 hrs
- Facturación total Q1: S/. 487,200.00
* Nuevos clientes captados: 1
"""


def test_heuristic_detecta_titulos_numerados():
    blocks = docx_service._heuristic_blocks(OCR_SAMPLE)
    headings = [b for b in blocks if b["type"] == "heading"]
    assert any("Resumen Ejecutivo" in h["text"] for h in headings)
    assert any("Métricas" in h["text"] for h in headings)


def test_heuristic_agrupa_vinetas():
    blocks = docx_service._heuristic_blocks(OCR_SAMPLE)
    bullets = [b for b in blocks if b["type"] == "bullets"]
    assert len(bullets) == 1
    items = bullets[0]["items"]
    assert "Total de horas facturadas: 1,840 hrs" in items
    assert len(items) == 3
    # El marcador (+, -, *) debe quedar removido del texto del ítem.
    assert not any(i.startswith(("+", "-", "*")) for i in items)


def test_heuristic_une_lineas_en_parrafo():
    blocks = docx_service._heuristic_blocks(OCR_SAMPLE)
    paras = [b for b in blocks if b["type"] == "paragraph"]
    assert any("Nexus\nConsulting" not in p["text"] and "Nexus Consulting" in p["text"]
               for p in paras)


def _read_docx(data: bytes):
    from docx import Document
    return Document(io.BytesIO(data))


def test_build_docx_desde_texto_plano():
    data = docx_service.build_docx(OCR_SAMPLE, "reporte.pdf")
    doc = _read_docx(data)
    textos = [p.text for p in doc.paragraphs]
    assert any("Documento digitalizado" in t for t in textos)
    assert any("Resumen Ejecutivo" in t for t in textos)


def test_build_docx_desde_bloques_con_tabla():
    blocks = [
        {"type": "heading", "level": 2, "text": "2. Estado de Proyectos"},
        {"type": "table",
         "header": ["Proyecto", "Avance"],
         "rows": [["DocMind IA", "68%"], ["Outsourcing Cloud", "Operativo"]]},
    ]
    data = docx_service.build_docx("texto ocr de respaldo", "reporte.pdf", blocks=blocks)
    doc = _read_docx(data)

    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert len(table.rows) == 3          # header + 2 filas
    assert table.rows[0].cells[0].text == "Proyecto"
    assert table.rows[1].cells[0].text == "DocMind IA"
    assert table.rows[2].cells[1].text == "Operativo"


def test_build_docx_vacio_pone_placeholder():
    data = docx_service.build_docx("", "vacio.pdf")
    doc = _read_docx(data)
    assert any("No se pudo extraer texto" in p.text for p in doc.paragraphs)
